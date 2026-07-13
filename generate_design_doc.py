#!/usr/bin/env python3
"""
生成设计说明书的docx文件
用法：python generate_design_doc.py <输入markdown文件> <输出docx文件> <文档标题>
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, cn="宋体", ascii_="Times New Roman", size=12, bold=False):
    """设置字体"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = ascii_
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), ascii_)
    rFonts.set(qn('w:hAnsi'), ascii_)


def setup_page(section):
    """设置页面"""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_heading(doc, text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = 16
        bold = True
        font_cn = "黑体"
    elif level == 2:
        size = 14
        bold = True
        font_cn = "黑体"
    elif level == 3:
        size = 12
        bold = True
        font_cn = "黑体"
    else:
        size = 12
        bold = True
        font_cn = "宋体"
    
    set_font(p.add_run(text), font_cn, "Arial", size, bold)
    return p


def add_body(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0.74)
    set_font(p.add_run(text))
    return p


def add_code_block(doc, text):
    """添加代码块"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, "Consolas", "Consolas", 10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_font(run, "黑体", "Arial", 11, True)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = ''
            p = row_cells[col_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, "宋体", "Times New Roman", 10.5)
    
    # 添加空行
    doc.add_paragraph()
    return table


def parse_markdown(md_text):
    """解析markdown，返回结构化的内容列表"""
    lines = md_text.split('\n')
    elements = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            elements.append(('heading', level, text))
            i += 1
            continue
        
        # 代码块
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束的```
            elements.append(('code', '\n'.join(code_lines)))
            continue
        
        # 表格
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # 解析表头
            headers = [h.strip() for h in line.split('|') if h.strip()]
            i += 2  # 跳过分隔线
            rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                rows.append(cells)
                i += 1
            elements.append(('table', headers, rows))
            continue
        
        # 空行
        if not line.strip():
            i += 1
            continue
        
        # 普通段落
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('```') and not ('|' in lines[i] and i + 1 < len(lines) and '---' in lines[i + 1]):
            para_lines.append(lines[i])
            i += 1
        text = ' '.join(para_lines)
        # 移除markdown格式（粗体、斜体等）
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        elements.append(('paragraph', text))
    
    return elements


def generate_docx(md_file, output_file, doc_title):
    """从markdown生成docx"""
    # 读取markdown文件
    md_text = Path(md_file).read_text(encoding='utf-8')
    
    # 解析markdown
    elements = parse_markdown(md_text)
    
    # 创建文档
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    setup_page(section)
    
    # 添加文档标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    set_font(title_p.add_run(doc_title), "黑体", "Arial", 22, True)
    
    # 添加内容
    for elem in elements:
        if elem[0] == 'heading':
            level = elem[1]
            text = elem[2]
            add_heading(doc, text, level)
        elif elem[0] == 'paragraph':
            text = elem[1]
            add_body(doc, text)
        elif elem[0] == 'code':
            text = elem[1]
            add_code_block(doc, text)
        elif elem[0] == 'table':
            headers = elem[1]
            rows = elem[2]
            add_table(doc, headers, rows)
    
    # 保存文档
    doc.save(output_file)
    kb = Path(output_file).stat().st_size / 1024
    print(f"已生成：{output_file}（{kb:.1f} KB）")


def main():
    if len(sys.argv) < 4:
        print("用法：python generate_design_doc.py <输入markdown文件> <输出docx文件> <文档标题>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2]
    doc_title = sys.argv[3]
    
    if not Path(md_file).exists():
        print(f"错误：文件不存在：{md_file}", file=sys.stderr)
        sys.exit(1)
    
    generate_docx(md_file, output_file, doc_title)


if __name__ == "__main__":
    main()
